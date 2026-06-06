"""
ByteRescue — Project Report Generator
Run this script to produce ByteRescue_Report.docx on the Desktop.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop", "ByteRescue_Report.docx")

# ─── Colour palette ───────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0B, 0x24, 0x48)   # headings
TEAL        = RGBColor(0x00, 0x7A, 0x87)   # sub-headings / accents
BLACK       = RGBColor(0x00, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xD6, 0xE4, 0xF0)   # table header bg
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)   # alternate row

# ─── Helper utilities ─────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=DARK_NAVY):
    p    = doc.add_heading(text, level=level)
    run  = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color
    run.font.bold      = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = TEAL
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = TEAL
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=BLACK, indent=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p


def add_bullet(doc, text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    return p


def add_code(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x60, 0x20)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "0B2448")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold      = True
            run.font.color.rgb = WHITE
            run.font.size      = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg    = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = val
            set_cell_bg(cells[ci], bg)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def page_break(doc):
    doc.add_page_break()


# ─── BUILD DOCUMENT ───────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ═══════════════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("ByteRescue")
    tr.font.size      = Pt(36)
    tr.font.bold      = True
    tr.font.color.rgb = DARK_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Signature-Based File Carving & Forensic Recovery Tool")
    sr.font.size      = Pt(16)
    sr.font.italic    = True
    sr.font.color.rgb = TEAL

    doc.add_paragraph()

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_p.add_run("A Python-Based Digital Forensics Project").font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [("Submitted By", "Harsh Belwal  &  Ekta Kaur"),
                          ("Subject", "Digital Forensics / Cyber Security"),
                          ("Tool Version", "ByteRescue v1.0"),
                          ("Language", "Python 3.x")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}:  ")
        r1.font.bold      = True
        r1.font.size      = Pt(12)
        r1.font.color.rgb = DARK_NAVY
        r2 = p.add_run(value)
        r2.font.size      = Pt(12)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    #  1. PROJECT OVERVIEW
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Project Overview", 1)
    add_para(doc,
        "ByteRescue is a custom-built, Python-based digital forensic file carving tool designed "
        "to recover deleted, lost, or inaccessible files from storage media — including USB drives, "
        "SD cards, and HDDs — even when the file system has been completely corrupted, formatted, or erased.")
    doc.add_paragraph()
    add_para(doc,
        "Unlike conventional recovery software that relies on the file system index (FAT32, NTFS, exFAT) "
        "to locate files, ByteRescue operates at the raw binary (byte) level. It reads the storage device "
        "sector-by-sector, identifies known file signatures (magic bytes), and reconstructs files without "
        "needing any file system metadata. This technique is known in Digital Forensics as File Carving "
        "or Raw Data Carving.")
    doc.add_paragraph()
    add_para(doc, "Key features include:", bold=True)
    add_bullet(doc, "Professional Military/OSINT-themed Graphical User Interface (GUI) built with Tkinter")
    add_bullet(doc, "Multi-threaded scanning for high-speed recovery")
    add_bullet(doc, "Automated JSON metadata logging for every recovered file (forensic-grade)")
    add_bullet(doc, "Exact byte-offset tracking for Hex Editor verification")
    add_bullet(doc, "Category-based, structured output folder system")

    # ═══════════════════════════════════════════════════════════════════════════
    #  2. PROBLEM STATEMENT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Problem Statement", 1)
    add_para(doc,
        "When a user accidentally deletes files, formats a drive, or when a drive gets corrupted, "
        "the actual data does not vanish immediately. The operating system only removes the file system's "
        "reference (index) to that data. The raw bytes of the original file remain on the disk until "
        "new data physically overwrites them.")
    doc.add_paragraph()
    add_para(doc,
        "The Challenge: How do you recover that data when the index pointing to it is gone? "
        "Standard tools fail here because they depend on the file system table. ByteRescue solves this "
        "by bypassing the file system entirely and scanning the raw binary data for known file patterns.",
        italic=True)

    # ═══════════════════════════════════════════════════════════════════════════
    #  3. OBJECTIVE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Objectives", 1)
    for obj in [
        "Build a tool that can carve (extract) files from raw disk images and physical storage devices.",
        "Support multiple file formats: images, documents, videos, compressed archives.",
        "Implement accurate, forensically-valid byte-offset tracking for each recovered file.",
        "Generate structured JSON metadata automatically during recovery for forensic reporting.",
        "Provide a professional, easy-to-use GUI that does not require command-line knowledge.",
    ]:
        add_bullet(doc, obj)

    # ═══════════════════════════════════════════════════════════════════════════
    #  4. TECHNOLOGY STACK
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Technology Stack & Libraries Used", 1)

    add_heading(doc, "4.1 Core Language", 2)
    make_table(doc,
        ["Technology", "Purpose"],
        [["Python 3.x", "Primary development language for the entire tool"]],
        col_widths=[2.5, 4.0])

    doc.add_paragraph()
    add_heading(doc, "4.2 Standard Library Modules", 2)
    make_table(doc,
        ["Module", "Purpose"],
        [
            ["tkinter",           "Building the Graphical User Interface (GUI)"],
            ["tkinter.ttk",       "Themed widgets: progress bars, spinboxes, dropdowns"],
            ["threading",         "Running multiple scanner threads in parallel"],
            ["struct",            "Parsing binary data (MP4 atoms, AVI RIFF headers)"],
            ["pathlib.Path",      "Cross-platform file system path management"],
            ["json",              "Writing and reading JSON metadata files"],
            ["ctypes",            "Low-level Windows API calls for raw disk access"],
            ["hashlib",           "MD5 hashing (used in earlier deduplication logic)"],
            ["os",                "OS-level file and environment operations"],
            ["time",              "Measuring and displaying total scan duration"],
            ["math",              "GUI animation: sine wave, hex grid calculations"],
            ["random",            "Particle animations in the GUI background"],
            ["shutil",            "Disk usage queries for logical fallback size detection"],
        ],
        col_widths=[1.9, 4.6])

    doc.add_paragraph()
    add_heading(doc, "4.3 External Tools Used", 2)
    make_table(doc,
        ["Tool", "Purpose"],
        [
            ["HxD Hex Editor",           "Independent forensic offset verification"],
            ["Windows Kernel32 API",     "Raw disk access via CreateFileW, DeviceIoControl"],
            ["PyInstaller",              "Packaging tool into standalone .exe (ByteRescue.spec)"],
        ],
        col_widths=[2.2, 4.3])

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    #  5. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. System Architecture", 1)

    add_para(doc, "Project Module Structure:", bold=True)
    for line in [
        "ByteRescue/",
        "  ├── gui.py              ← Main GUI entry point (Military OSINT theme)",
        "  ├── main.py             ← CLI entry point & scan orchestration",
        "  ├── config.py           ← Global settings, file signatures dictionary",
        "  ├── signatures.py       ← Active signature loader",
        "  ├── scanner.py          ← Core carving engine (RecoveryThread class)",
        "  ├── mp4_recovery.py     ← Specialised MP4 atom-walking recovery",
        "  ├── drive_utils.py      ← Windows RAW disk access & drive enumeration",
        "  ├── utils.py            ← Output structure, file naming, helpers",
        "  ├── forensic_logger.py  ← JSON metadata logging system",
        "  ├── progress.py         ← Live progress monitor",
        "  └── report.py           ← Text & JSON scan report generation",
    ]:
        add_code(doc, line)

    doc.add_paragraph()
    add_para(doc, "Execution Flow:", bold=True)
    for step in [
        "1. User selects drive in GUI  ->  drive_utils.py opens \\\\.\\E: as raw binary",
        "2. Drive size is detected  →  main.py splits into N equal chunks",
        "3. N RecoveryThreads launched by scanner.py — each scans its range",
        "4. scan_chunk() calls chunk.find(signature) for all known file types",
        "5. is_valid_signature() validates each hit against type-specific rules",
        "6. recover_file() dispatches to the correct carving strategy",
        "7. forensic_logger.py appends byte-offset entry to the JSON file",
        "8. All threads join → report.py writes scan summary",
    ]:
        add_bullet(doc, step)

    # ═══════════════════════════════════════════════════════════════════════════
    #  6. CORE FORENSIC CONCEPTS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Core Forensic Concepts", 1)

    add_heading(doc, "6.1 File Carving", 2)
    add_para(doc,
        "File carving is the process of extracting files from a raw binary stream without relying on any "
        "file system metadata. It works because all files of a given format share a known, standardised "
        "header (and often a footer) byte sequence, regardless of what file system they are stored in.")

    add_heading(doc, "6.2 File Signatures (Magic Numbers)", 2)
    add_para(doc,
        "Every file format carries a unique sequence of bytes at its beginning called a magic number or "
        "file signature. These are internationally standardised — the same bytes appear in every valid "
        "file of that type, on any operating system, in any country.")

    add_heading(doc, "6.3 Sectors & Byte Offsets", 2)
    add_para(doc,
        "A storage device is divided into sectors (512 bytes each). Files are stored in multiples of "
        "clusters. ByteRescue tracks two values for every discovered file:")
    add_bullet(doc, "Byte Offset — The exact byte position on the device (e.g., 5,537,792)")
    add_bullet(doc, "Hex Offset — Same in hexadecimal (e.g., 0x548000) — used in Hex Editor navigation")
    add_bullet(doc, "Sector Number — Byte_Offset / 512 = 10,816 — for cross-referencing")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    #  7. COMPLETE FILE SIGNATURES TABLE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Complete File Signatures Used in ByteRescue", 1)
    add_para(doc,
        "The following table lists every file signature implemented in the ByteRescue engine. "
        "The hex values are the actual bytes that the scanner searches for in the raw binary stream. "
        "These are the same values visible in any Hex Editor when you open a valid file of that type.")

    doc.add_paragraph()

    sig_rows = [
        # IMAGES
        ("Images",      "JPEG",   ".jpg / .jpeg", "FF D8 FF",                       "FF D9",            "2",    "Footer"),
        ("Images",      "PNG",    ".png",          "89 50 4E 47 0D 0A 1A 0A",       "IEND AE 42 60 82", "8",    "Footer"),

        # DOCUMENTS
        ("Documents",   "PDF",    ".pdf",          "25 50 44 46  (%PDF)",            "25 25 45 4F 46 (%%EOF)", "1024", "Footer"),
        ("Documents",   "DOCX",   ".docx",         "50 4B 03 04  (PK..)",            "50 4B 05 06",      "22",   "Footer"),
        ("Documents",   "DOC",    ".doc",          "D0 CF 11 E0 A1 B1 1A E1",        "None",             "—",    "Heuristic"),
        ("Documents",   "XLS",    ".xls",          "D0 CF 11 E0 A1 B1 1A E1",        "None",             "—",    "Heuristic"),
        ("Documents",   "PPT",    ".ppt",          "D0 CF 11 E0 A1 B1 1A E1",        "None",             "—",    "Heuristic"),

        # VIDEOS
        ("Videos",      "MP4",    ".mp4",          "66 74 79 70  (ftyp)",            "None",             "—",    "Atom Walk"),
        ("Videos",      "AVI",    ".avi",          "52 49 46 46 .. .. .. .. 41 56 49 20  (RIFF..AVI )", "None", "—", "RIFF Header"),
        ("Videos",      "MKV",    ".mkv",          "1A 45 DF A3",                    "None",             "—",    "Heuristic"),
        ("Videos",      "MOV",    ".mov",          "6D 6F 6F 76  (moov)",            "None",             "—",    "Heuristic"),
        ("Videos",      "FLV",    ".flv",          "46 4C 56  (FLV)",                "None",             "—",    "Heuristic"),

        # COMPRESSED
        ("Compressed",  "ZIP",    ".zip",          "50 4B 03 04  (PK..)",            "50 4B 05 06",      "22",   "Footer"),
        ("Compressed",  "RAR",    ".rar",          "52 61 72 21 1A 07  (Rar!..)",    "None",             "—",    "Heuristic"),
        ("Compressed",  "7-Zip",  ".7z",           "37 7A BC AF 27 1C  (7z...)",     "None",             "—",    "Heuristic"),
        ("Compressed",  "GZIP",   ".gz",           "1F 8B 08",                       "None",             "—",    "Heuristic"),
        ("Compressed",  "ISO",    ".iso",          "43 44 30 30 31  (CD001)",        "None",             "—",    "Heuristic"),

        # AUDIO
        ("Audio",       "MP3",    ".mp3",          "49 44 33  (ID3)",                "None",             "—",    "Heuristic"),

        # MISC
        ("Misc",        "TIFF",   ".tiff",         "49 49 2A 00 / 4D 4D 00 2A",     "None",             "—",    "Heuristic"),
        ("Misc",        "PST",    ".pst",          "21 42 44 4E  (!BDN)",            "None",             "—",    "Heuristic"),
    ]

    make_table(doc,
        ["Category", "Format", "Extension", "Header Signature (Hex)", "Footer Signature", "End\nOffset", "Recovery\nMethod"],
        sig_rows,
        col_widths=[0.9, 0.7, 0.9, 2.3, 1.5, 0.6, 0.9])

    doc.add_paragraph()
    add_para(doc,
        "Notes on Recovery Methods:", bold=True)
    add_bullet(doc, "Footer — Tool reads bytes from header until the known footer sequence is found, then saves exactly those bytes.")
    add_bullet(doc, "Atom Walk — For MP4: the ISO Box structure (ftyp, moov, mdat atoms) is walked; their declared sizes are added to determine the file end.")
    add_bullet(doc, "RIFF Header — For AVI: bytes 4–8 of the RIFF header declare the payload size; tool reads exactly that many bytes.")
    add_bullet(doc, "Heuristic — Tool reads forward until the next different file signature appears at a sector boundary (512-byte aligned).")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    #  8. MODULE BREAKDOWN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Module-by-Module Breakdown", 1)

    modules = [
        (
            "config.py — Global Configuration",
            "Contains all tunable parameters and the FILE_SIGNATURES dictionary — the single "
            "source of truth for what file types are detected and how.",
            [
                "SCAN_CHUNK_SIZE = 1 MB   — Read size per iteration",
                "MAX_RECOVERY_FILE_SIZE = 200 MB  — Safety cap for documents",
                "MAX_VIDEO_SIZE = 500 MB         — Safety cap for videos",
                "DEFAULT_THREAD_COUNT = 4        — Parallel scanner threads",
            ]
        ),
        (
            "scanner.py — Core Carving Engine",
            "Contains the RecoveryThread class. Each thread opens the raw disk, reads 1MB chunks "
            "with 2MB overlap at boundaries, validates signatures, and dispatches to the correct "
            "recovery strategy. Thread-safe offset deduplication prevents double-recovery.",
            []
        ),
        (
            "mp4_recovery.py — MP4 Atom Walker",
            "MP4 files have no footer. This module walks the ISO Base Media Box (Atom) structure. "
            "Each atom has a 4-byte big-endian size + 4-byte name. Valid atoms (ftyp, moov, mdat, "
            "trak, mdia, stbl…) are accumulated until an invalid atom is found, giving the exact file end.",
            []
        ),
        (
            "forensic_logger.py — JSON Metadata System",
            "For every successfully carved file: determines category, generates sequential filename "
            "(image1.jpg), records exact byte offsets in decimal and hex, and appends to the "
            "per-extension JSON file (jpg.json, pdf.json). Fully thread-safe via threading.Lock().",
            []
        ),
        (
            "drive_utils.py — Raw Disk Access",
            "Uses ctypes to call Windows Kernel32 CreateFileW with GENERIC_READ on \\\\.\\E: to get "
            "a raw handle bypassing the file system. Uses IOCTL_DISK_GET_LENGTH_INFO to get exact "
            "drive size in bytes. Falls back to shutil.disk_usage for logical fallback.",
            []
        ),
        (
            "gui.py — Graphical User Interface",
            "A full Tkinter GUI with Military OSINT / Classified Terminal aesthetic. Features: "
            "animated hex grid background, rotating radar sweep, circular arc progress meter, "
            "colour-coded live log terminal, and a recovery summary card. Completely decoupled "
            "from the scan engine via threading callbacks.",
            []
        ),
    ]

    for title, desc, bullets in modules:
        add_heading(doc, title, 2)
        add_para(doc, desc)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    #  9. GUI OVERVIEW
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. GUI Overview", 1)

    sections_gui = [
        ("Drive Selection Panel",
         "Dropdown list of available logical drives. Drive size is auto-detected. Thread count adjusted via spinbox."),
        ("Animated Radar Display",
         "Rotating radar sweep animation with animated hex grid background to create a live-system feel."),
        ("Scan Control",
         "BEGIN SCAN button starts multi-threaded scan. Circular arc indicator shows % of drive scanned. "
         "File counter updates in real time with neon glow effect."),
        ("Live Terminal Log",
         "Scrolling log with entries colour-coded by file type: Orange (images), Teal (videos), Red (PDF), Purple (ZIP)."),
        ("Recovery Summary Card",
         "After scan completion, a breakdown card shows count recovered per file type and total time taken."),
    ]
    for panel, desc in sections_gui:
        add_para(doc, panel + ":", bold=True)
        add_para(doc, desc, indent=True)
        doc.add_paragraph()

    add_para(doc,
        "[ GUI Screenshot — Insert screenshot of the ByteRescue application window here ]",
        italic=True, color=TEAL)
    doc.add_paragraph()
    add_para(doc,
        "[ Output Screenshot — Insert screenshot of the Recovered_Files folder structure here ]",
        italic=True, color=TEAL)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    #  10. OUTPUT STRUCTURE & JSON
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Output Structure & JSON Metadata", 1)

    add_para(doc, "Folder Structure Generated After Scan:", bold=True)
    for line in [
        "Recovered_Files/",
        "  ├── images/",
        "  │   ├── jpg/",
        "  │   │   ├── image1.jpg",
        "  │   │   ├── image2.jpg",
        "  │   │   └── jpg.json     ← JSON log for all recovered JPGs",
        "  │   └── png/",
        "  │       ├── image1.png",
        "  │       └── png.json",
        "  ├── documents/",
        "  │   ├── pdf/",
        "  │   │   ├── document1.pdf",
        "  │   │   └── pdf.json",
        "  │   └── docx/",
        "  │       ├── document1.docx",
        "  │       └── docx.json",
        "  ├── videos/",
        "  │   ├── mp4/",
        "  │   │   ├── video1.mp4",
        "  │   │   └── mp4.json",
        "  └── reports/",
        "      ├── scan_report_YYYY-MM-DD.txt",
        "      └── scan_report_YYYY-MM-DD.json",
    ]:
        add_code(doc, line)

    doc.add_paragraph()
    add_para(doc, "Sample JSON Entry (images/jpg/jpg.json):", bold=True)
    for line in [
        "[",
        "  {",
        '    "file_name": "image1.jpg",',
        '    "exact_start_hex_offset": "0x548000",',
        '    "exact_start_byte_offset": 5537792,',
        '    "exact_end_hex_offset": "0x923216",',
        '    "exact_end_byte_offset": 9581078,',
        '    "approx_sector_number": 10816',
        "  },",
        "  { ... }",
        "]",
    ]:
        add_code(doc, line)

    # ═══════════════════════════════════════════════════════════════════════════
    #  11. STEP BY STEP
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. How the Tool Works — Step by Step", 1)

    steps = [
        ("Step 1: Launch & Drive Selection",
         "User opens ByteRescue GUI and selects their USB drive from the dropdown. Drive size is auto-detected using ctypes."),
        ("Step 2: Raw Device Open",
         "drive_utils.py uses Windows Kernel32 API to open \\\\.\\E: as raw binary - bypassing FAT32/NTFS completely."),
        ("Step 3: Drive Partitioning",
         "Total size is divided by thread count (default 4). Each thread gets equal range with 2MB overlap at boundaries."),
        ("Step 4: Signature Search",
         "Each thread calls chunk.find(signature) for every known type in each 1MB block — extremely fast (C-level performance)."),
        ("Step 5: Signature Validation",
         "is_valid_signature() confirms the hit is genuine: checks byte after FF D8 for JPEG, box size for MP4, version byte for ZIP."),
        ("Step 6: File Carving",
         "Correct strategy runs: Footer-based for JPG/PDF, Atom Walk for MP4, RIFF Header for AVI, Heuristic for others."),
        ("Step 7: JSON Logging",
         "forensic_logger.py immediately appends entry to jpg.json with exact byte offset — thread-safe, no duplicates."),
        ("Step 8: Scan Complete",
         "All threads join. report.py writes text and JSON summary reports. GUI shows recovery count per file type."),
    ]

    for title, desc in steps:
        add_para(doc, title + ":", bold=True)
        add_para(doc, desc, indent=True)

    # ═══════════════════════════════════════════════════════════════════════════
    #  12. HEX EDITOR VERIFICATION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Forensic Validation — Hex Editor Verification", 1)
    add_para(doc,
        "Every byte offset recorded in the JSON files can be independently verified using HxD Hex Editor "
        "or any equivalent tool. This is standard forensic chain-of-custody validation.")
    doc.add_paragraph()
    add_para(doc, "Verification Procedure:", bold=True)
    add_bullet(doc, "Open HxD Hex Editor")
    add_bullet(doc, "Go to File → Open Disk → Logical Disks → Select your drive (E:)")
    add_bullet(doc, "Press Ctrl+G (Go To Offset)")
    add_bullet(doc, "Select Hex mode and type the exact_start_hex_offset from JSON (e.g., 548000)")
    add_bullet(doc, "Cursor will land exactly on the first byte of the file signature (e.g., FF D8 FF E0 for JPEG)")
    doc.add_paragraph()
    add_para(doc,
        "Important: Always use exact_start_hex_offset for navigation. "
        "The approx_sector_number is for approximate reference only — Hex Editor navigation "
        "requires the exact byte offset for pixel-perfect precision.",
        italic=True, color=TEAL)

    # ═══════════════════════════════════════════════════════════════════════════
    #  13. RESULTS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13. Results & Testing", 1)
    add_para(doc,
        "The tool was tested on a 16GB USB Pendrive that had been Quick-Formatted (FAT32), "
        "meaning its file system index was erased but raw data remained.")
    doc.add_paragraph()

    make_table(doc,
        ["File Type", "Files Present", "Files Recovered", "Success Rate", "Method Used"],
        [
            ["JPEG",  "~30",  "~27",  "~90%",  "Footer (FF D9)"],
            ["PNG",   "~5",   "~5",   "100%",  "Footer (IEND)"],
            ["PDF",   "~6",   "~6",   "100%",  "Footer (%%EOF)"],
            ["MP4",   "~3",   "~2",   "~67%",  "Atom Walk"],
            ["DOCX",  "~4",   "~3",   "~75%",  "Footer (PK 05 06)"],
        ],
        col_widths=[1.2, 1.2, 1.4, 1.2, 2.5])

    doc.add_paragraph()
    add_para(doc,
        "JSON Offset Verification: Every recovered file's exact_start_hex_offset was independently "
        "verified in HxD — all offsets matched the file's magic number exactly, confirming 100% "
        "offset accuracy.", bold=True)
    doc.add_paragraph()
    add_para(doc, "Scan Speed: 16GB drive scan completed in approximately 8–12 minutes using 4 threads.")

    # ═══════════════════════════════════════════════════════════════════════════
    #  14. CHALLENGES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14. Challenges Faced & Solutions", 1)

    make_table(doc,
        ["Challenge", "Solution"],
        [
            ["MP4 files have no footer — difficult to determine file end",
             "Implemented ISO Base Media Atom Walker in mp4_recovery.py"],
            ["DOCX and ZIP share the same PK\\x03\\x04 magic bytes",
             "Added ZIP version byte validation (ver <= 63) in is_valid_signature()"],
            ["Signature found mid-chunk; real MP4 file starts 4 bytes earlier",
             "Implemented RECOVERY_REWIND dictionary to seek back correctly"],
            ["Multiple threads recovering the same file at the same offset",
             "Used saved_offsets set with threading.Lock() for atomic deduplication"],
            ["JSON writing from multiple threads causing race conditions",
             "Implemented _log_lock for atomic JSON read → append → write cycle"],
            ["%%EOF in PDF appears multiple times in complex/linearised PDFs",
             "Added 1024-byte end_offset headroom to capture full cross-reference table"],
        ],
        col_widths=[3.1, 3.4])

    # ═══════════════════════════════════════════════════════════════════════════
    #  15. FUTURE SCOPE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "15. Future Scope", 1)
    scope = [
        ("AI Integration",        "Feed the JSON metadata to an LLM API (Gemini/GPT-4) for automated natural-language forensic analysis reports."),
        ("Disk Image Support",    "Allow scanning of .dd, .img, and .E01 forensic image files instead of live physical disks."),
        ("More File Types",       "Add signatures for RAW camera formats (.CR2, .NEF), SQLite databases, email archives (.eml, .mbox)."),
        ("Timeline View",         "Plot recovered files on a graphical timeline based on their sector position on the disk."),
        ("Integrity Hashing",     "Generate SHA-256 hash of each carved file and log it in JSON for chain-of-custody evidence."),
        ("Cross-Platform",        "Extend Linux/macOS support using /dev/sdX raw device access instead of Windows Kernel32 API."),
    ]
    for title, desc in scope:
        add_para(doc, f"{title}: ", bold=True)
        add_para(doc, desc, indent=True)

    # ═══════════════════════════════════════════════════════════════════════════
    #  16. CONCLUSION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "16. Conclusion", 1)
    add_para(doc,
        "ByteRescue demonstrates the practical application of Digital Forensics principles in a "
        "real Python-based tool. By bypassing the file system and operating directly at the byte level, "
        "it successfully recovers files even from formatted or corrupted storage media.")
    doc.add_paragraph()
    add_para(doc,
        "The key innovation beyond basic carving is the automated, structured JSON metadata logging "
        "system integrated into the recovery pipeline. This provides investigators with precise "
        "byte-level traceability, file-type organised evidence folders, and independently verifiable "
        "forensic offsets.")
    doc.add_paragraph()
    add_para(doc,
        "The tool was built entirely in Python using the standard library, demonstrating that "
        "sophisticated forensic tools do not require heavy commercial frameworks. ByteRescue is "
        "structured, modular, documented, and ready for use as a practical forensic aid in "
        "educational and investigative contexts.")

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Report —")
    r.font.italic    = True
    r.font.color.rgb = TEAL
    r.font.size      = Pt(12)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Harsh Belwal  &  Ekta Kaur  |  ByteRescue v1.0")
    r2.font.size      = Pt(10)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"\nReport saved to:\n    {OUTPUT_PATH}\n")


if __name__ == "__main__":
    build()
